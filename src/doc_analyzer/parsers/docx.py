"""Word (.docx) parsing via python-docx."""

from __future__ import annotations

import logging
from pathlib import Path

from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError

from doc_analyzer.models import Document

logger = logging.getLogger(__name__)


class DocxParserError(Exception):
    """Base class for DOCX parsing failures."""


class DocxFileNotFoundError(DocxParserError):
    """The path does not exist or is not a regular file."""


class DocxMalformedError(DocxParserError):
    """python-docx cannot parse the file as a valid .docx."""


def parse_docx(path: str | Path) -> Document:
    """Parse a Word (.docx) file into a :class:`Document`.

    DOCX files have no inherent page concept, so the whole document is
    returned as a single page made of paragraphs joined by newlines, plus
    table content extracted row-by-row.
    """
    docx_path = Path(path)
    if not docx_path.is_file():
        raise DocxFileNotFoundError(f"DOCX file does not exist: {docx_path}")

    try:
        doc = DocxDocument(str(docx_path))
    except PackageNotFoundError as exc:
        raise DocxMalformedError(f"Failed to open DOCX: {docx_path}") from exc
    except Exception as exc:
        raise DocxMalformedError(f"Unexpected error reading DOCX: {docx_path}: {exc}") from exc

    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip("| "):
                parts.append(row_text)

    page_text = "\n".join(parts) if parts else ""

    metadata: dict[str, str] = {}
    try:
        core = doc.core_properties
        if core.title:
            metadata["title"] = str(core.title)
        if core.author:
            metadata["author"] = str(core.author)
        if core.subject:
            metadata["subject"] = str(core.subject)
    except Exception as exc:
        logger.debug("Could not read DOCX core properties: %s", exc)

    return Document(source=docx_path, pages=(page_text,), metadata=metadata)
